"""Write filled templates to professional PDF, .docx, or .md output files.

PDF uses reportlab for executive-grade, investor-quality formatting with
branded letterhead, logo, and corporate color scheme.
"""
from __future__ import annotations

import re
import datetime
from pathlib import Path


def _strip_md(s: str) -> str:
    return re.sub(r"\*\*(.+?)\*\*", r"\1", s)

def _safe(s: str) -> str:
    return (s.replace("\u2022", "-").replace("\u2610", "[ ]")
             .replace("\u2013", "-").replace("\u2014", "--")
             .replace("\u2018", "'").replace("\u2019", "'")
             .replace("\u201c", '"').replace("\u201d", '"')
             .replace("\u2026", "..."))


def write_markdown(text: str, output_path: Path) -> Path:
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(text, encoding="utf-8")
    return output_path


def write_docx(text: str, output_path: Path, title: str = "Document") -> Path:
    from docx import Document
    from docx.shared import Pt, Inches
    doc = Document()
    style = doc.styles["Normal"]; style.font.name = "Calibri"; style.font.size = Pt(11)
    for s in doc.sections: s.top_margin = s.bottom_margin = Inches(1); s.left_margin = s.right_margin = Inches(1)
    for line in text.split("\n"):
        s = line.strip()
        if not s: doc.add_paragraph("")
        elif s.startswith("# "): doc.add_heading(s[2:], level=1)
        elif s.startswith("## "): doc.add_heading(s[3:], level=2)
        elif s.startswith("### "): doc.add_heading(s[4:], level=3)
        elif s.startswith("---"): doc.add_paragraph("").paragraph_format.space_after = Pt(6)
        else:
            p = doc.add_paragraph()
            for part in re.split(r"(\*\*.*?\*\*)", s):
                if part.startswith("**") and part.endswith("**"): p.add_run(part[2:-2]).bold = True
                else: p.add_run(part)
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


def write_pdf(text: str, output_path: Path, title: str = "Document",
              company: str = "", phone: str = "", email: str = "",
              address: str = "", website: str = "", tagline: str = "",
              logo_path: str = "", property_name: str = "") -> Path:
    """Generate an executive-grade branded PDF."""
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.units import inch
    from reportlab.lib.colors import HexColor
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.enums import TA_LEFT, TA_RIGHT
    from reportlab.platypus import (
        BaseDocTemplate, PageTemplate, Frame,
        Paragraph, Spacer, Table, TableStyle, HRFlowable, Image,
    )

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    W, H = letter
    M = 0.7 * inch
    navy = HexColor("#1B365D")
    gold = HexColor("#C5922E")
    dk = HexColor("#2D3748")
    md = HexColor("#6B7280")
    lt = HexColor("#F3F4F6")
    wh = HexColor("#FFFFFF")

    # Check if logo exists.
    has_logo = logo_path and Path(logo_path).exists()
    HEADER_H = 80 if has_logo else 65
    now = datetime.datetime.now().strftime("%B %d, %Y")

    # Styles
    sH1 = ParagraphStyle("H1", fontName="Helvetica-Bold", fontSize=15, textColor=navy, spaceBefore=16, spaceAfter=4, leading=19)
    sH2 = ParagraphStyle("H2", fontName="Helvetica-Bold", fontSize=12, textColor=navy, spaceBefore=12, spaceAfter=3, leading=15)
    sH3 = ParagraphStyle("H3", fontName="Helvetica-Bold", fontSize=10, textColor=dk, spaceBefore=8, spaceAfter=2, leading=13)
    sBody = ParagraphStyle("Body", fontName="Helvetica", fontSize=9.5, textColor=dk, spaceAfter=5, leading=13)
    sBullet = ParagraphStyle("Bullet", fontName="Helvetica", fontSize=9.5, textColor=dk, spaceAfter=2, leading=12, leftIndent=16, bulletIndent=4)
    sTH = ParagraphStyle("TH", fontName="Helvetica-Bold", fontSize=8.5, textColor=wh)
    sTC = ParagraphStyle("TC", fontName="Helvetica", fontSize=8.5, textColor=dk)

    def _header_footer(canvas, doc):
        canvas.saveState()
        # ── HEADER ──
        # Navy background
        canvas.setFillColor(navy)
        canvas.rect(0, H - HEADER_H, W, HEADER_H, fill=1, stroke=0)
        # Gold accent stripe
        canvas.setFillColor(gold)
        canvas.rect(0, H - HEADER_H - 3, W, 3, fill=1, stroke=0)

        # Logo
        if has_logo:
            try:
                canvas.drawImage(logo_path, M, H - HEADER_H + 8,
                                 width=HEADER_H - 16, height=HEADER_H - 16,
                                 preserveAspectRatio=True, mask='auto')
                text_x = M + HEADER_H - 6
            except Exception:
                text_x = M
        else:
            text_x = M

        # Company name
        canvas.setFillColor(wh)
        canvas.setFont("Helvetica-Bold", 16)
        canvas.drawString(text_x, H - 35, _safe(company or title))
        # Tagline
        if tagline:
            canvas.setFont("Helvetica", 8)
            canvas.setFillColor(HexColor("#94A3B8"))
            canvas.drawString(text_x, H - 48, _safe(tagline))
        # Website
        if website:
            canvas.setFont("Helvetica", 7.5)
            canvas.setFillColor(gold)
            canvas.drawString(text_x, H - 58, _safe(website))

        # Right side
        canvas.setFillColor(HexColor("#CBD5E1"))
        canvas.setFont("Helvetica", 8.5)
        canvas.drawRightString(W - M, H - 30, now)
        if property_name:
            canvas.setFont("Helvetica-Bold", 9)
            canvas.setFillColor(gold)
            canvas.drawRightString(W - M, H - 44, _safe(property_name))
        if phone:
            canvas.setFont("Helvetica", 8)
            canvas.setFillColor(HexColor("#94A3B8"))
            canvas.drawRightString(W - M, H - 56, _safe(phone))

        # ── FOOTER ──
        canvas.setStrokeColor(gold)
        canvas.setLineWidth(1.5)
        canvas.line(M, 45, W - M, 45)

        # Left: contact
        canvas.setFont("Helvetica", 7)
        canvas.setFillColor(md)
        footer_parts = [p for p in [phone, email, website] if p]
        canvas.drawString(M, 34, _safe(" | ".join(footer_parts)))
        if address:
            canvas.drawString(M, 24, _safe(address))

        # Right: page + company
        canvas.setFont("Helvetica", 7.5)
        canvas.drawRightString(W - M, 34, f"Page {doc.page}")
        canvas.setFont("Helvetica-Bold", 7)
        canvas.setFillColor(navy)
        canvas.drawRightString(W - M, 24, _safe(company or "AI Search"))

        canvas.restoreState()

    # Document frame
    frame = Frame(M, 55, W - 2*M, H - HEADER_H - 70,
                  topPadding=8, bottomPadding=8)
    pt = PageTemplate("main", frames=[frame], onPage=_header_footer)
    doc = BaseDocTemplate(str(output_path), pagesize=letter, pageTemplates=[pt],
                          leftMargin=M, rightMargin=M, topMargin=HEADER_H+10, bottomMargin=60)

    story = []

    def _rich(text, style):
        safe = _safe(text).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        safe = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", safe)
        safe = safe.replace("&lt;b&gt;", "<b>").replace("&lt;/b&gt;", "</b>")
        return Paragraph(safe, style)

    table_rows = []
    in_table = False

    for line in text.split("\n"):
        s = line.strip()
        if in_table and (not s.startswith("|") or not s.endswith("|")):
            if table_rows: _flush_table(story, table_rows, sTH, sTC, navy, gold, lt, dk)
            table_rows = []; in_table = False
        if not s: story.append(Spacer(1, 5)); continue
        if s.startswith("---"):
            story.append(HRFlowable(width="100%", color=gold, thickness=1.5, spaceAfter=6, spaceBefore=3)); continue
        if s.startswith("# "):
            story.append(_rich(s[2:], sH1))
            story.append(HRFlowable(width="35%", color=gold, thickness=2, spaceAfter=6)); continue
        if s.startswith("## "):
            story.append(_rich(s[3:], sH2))
            story.append(HRFlowable(width="20%", color=gold, thickness=1, spaceAfter=4)); continue
        if s.startswith("### "): story.append(_rich(s[4:], sH3)); continue
        if s.startswith("|") and s.endswith("|"):
            cells = [c.strip() for c in s.split("|")[1:-1]]
            if all(set(c) <= {"-", " ", ":"} for c in cells): continue
            table_rows.append(cells); in_table = True; continue
        if s.startswith("- [ ] "):
            story.append(Paragraph("[ ]  " + _safe(_strip_md(s[6:])), sBullet)); continue
        if s.startswith("- "):
            safe = _safe(s[2:])
            safe = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", safe)
            story.append(Paragraph("-  " + safe, sBullet)); continue
        story.append(_rich(s, sBody))

    if table_rows: _flush_table(story, table_rows, sTH, sTC, navy, gold, lt, dk)

    story.append(Spacer(1, 16))
    sfn = ParagraphStyle("fn", fontName="Helvetica-Oblique", fontSize=7, textColor=md, alignment=TA_RIGHT)
    story.append(Paragraph(f"Generated {now} | {_safe(company or 'AI Search')}", sfn))

    doc.build(story)
    return output_path


def _flush_table(story, rows, sH, sC, navy, gold, lt, dk):
    from reportlab.platypus import Table, TableStyle, Spacer, Paragraph
    from reportlab.lib.colors import HexColor
    if not rows: return
    data = []
    for i, row in enumerate(rows):
        data.append([Paragraph(_safe(c), sH if i == 0 else sC) for c in row])
    ncols = max(len(r) for r in data)
    col_w = 480 / ncols
    t = Table(data, colWidths=[col_w] * ncols)
    style = [
        ("BACKGROUND", (0,0), (-1,0), navy),
        ("TEXTCOLOR", (0,0), (-1,0), HexColor("#FFFFFF")),
        ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold"),
        ("FONTSIZE", (0,0), (-1,-1), 8.5),
        ("ALIGN", (0,0), (-1,-1), "LEFT"),
        ("VALIGN", (0,0), (-1,-1), "MIDDLE"),
        ("TOPPADDING", (0,0), (-1,-1), 4),
        ("BOTTOMPADDING", (0,0), (-1,-1), 4),
        ("LEFTPADDING", (0,0), (-1,-1), 6),
        ("RIGHTPADDING", (0,0), (-1,-1), 6),
        ("GRID", (0,0), (-1,-1), 0.5, HexColor("#D1D5DB")),
    ]
    for i in range(1, len(data)):
        if i % 2 == 0: style.append(("BACKGROUND", (0,i), (-1,i), lt))
    t.setStyle(TableStyle(style))
    story.append(Spacer(1, 4))
    story.append(t)
    story.append(Spacer(1, 6))


def fill_docx_template(template_path, replacements, output_path):
    from docx import Document
    doc = Document(str(template_path))
    def rip(para):
        full = "".join(run.text for run in para.runs)
        if "{{" not in full: return
        for k, v in replacements.items(): full = full.replace("{{"+k+"}}", v)
        full = re.sub(r"\{\{.+?\}\}", "[UNANSWERED]", full)
        if para.runs: para.runs[0].text = full
        for run in para.runs[1:]: run.text = ""
    for p in doc.paragraphs: rip(p)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs: rip(p)
    output_path = Path(output_path); output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path)); return output_path
